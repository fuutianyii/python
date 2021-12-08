from docx import Document
document = Document('F:\\Desktop\\社团活动(1).docx')  #打开文档
for tb1 in range(0,len(document.tables)):  #获取文档的表格个数 len(document.tables)
    tt=document.tables[tb1].rows  #按行读取第tb1个表的内容
    for low in tt:                #一行行读取tb1个表的内容
        row_cells=low.cells       #low.cells将行转换为列表
        if row_cells[0].text=='活动时间': #row_cells[0].text将行的第一个元素转换为文字
            row_cells[1].text='1'   #row_cells[1].text='1'将row_cells[1].text的元素替换成1
            document.save('complet.docx')  #另存为